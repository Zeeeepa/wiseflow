"""
Reference content extraction for Wiseflow.

This module provides functionality for extracting content from various file types
to be used as references for focus points.
"""

import os
import logging
import mimetypes
import tempfile
import contextlib
from typing import Dict, Any, Optional, Tuple
import requests
from urllib.parse import urlparse
import PyPDF2
import docx
import csv
import json
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import chardet

logger = logging.getLogger(__name__)

class ReferenceExtractor:
    """Extracts content from various file types for reference materials."""
    
    def __init__(self):
        """Initialize the reference extractor."""
        # Register additional MIME types
        mimetypes.add_type('application/pdf', '.pdf')
        mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
        mimetypes.add_type('text/csv', '.csv')
        mimetypes.add_type('application/json', '.json')
        mimetypes.add_type('application/xml', '.xml')
        mimetypes.add_type('text/html', '.html')
        mimetypes.add_type('text/html', '.htm')
    
    def extract_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract content from a file based on its type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple containing extracted text content and metadata
        """
        try:
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if not mime_type:
                # Try to determine type by reading file header
                mime_type = self._detect_mime_type(file_path)
            
            if mime_type:
                logger.debug(f"Detected MIME type: {mime_type} for file: {file_path}")
                
                # Extract content based on MIME type
                if mime_type == 'application/pdf':
                    return self._extract_pdf(file_path)
                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    return self._extract_docx(file_path)
                elif mime_type == 'text/csv':
                    return self._extract_csv(file_path)
                elif mime_type == 'application/json':
                    return self._extract_json(file_path)
                elif mime_type == 'application/xml' or mime_type == 'text/xml':
                    return self._extract_xml(file_path)
                elif mime_type.startswith('text/html'):
                    return self._extract_html(file_path)
                elif mime_type.startswith('text/'):
                    return self._extract_text(file_path)
            
            # Default to basic text extraction
            return self._extract_text(file_path)
        except Exception as e:
            logger.error(f"Error extracting content from file {file_path}: {e}")
            return "", {"error": str(e)}
    
    def extract_web_content(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract content from a web URL.
        
        Args:
            url: Web URL to extract content from
            
        Returns:
            Tuple containing extracted text content and metadata
        """
        try:
            with contextlib.suppress(requests.exceptions.RequestException):
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                
                # Create a temporary file to store the content
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    temp_file.write(response.content)
                    temp_path = temp_file.name
                
                try:
                    # Determine content type
                    content_type = response.headers.get('Content-Type', '').split(';')[0]
                    
                    metadata = {
                        "url": url,
                        "domain": urlparse(url).netloc,
                        "content_type": content_type,
                        "status_code": response.status_code,
                        "headers": dict(response.headers)
                    }
                    
                    # Extract content based on content type
                    if content_type == 'application/pdf':
                        content, pdf_metadata = self._extract_pdf(temp_path)
                        metadata.update(pdf_metadata)
                        return content, metadata
                    elif content_type == 'application/json':
                        content, json_metadata = self._extract_json(temp_path)
                        metadata.update(json_metadata)
                        return content, metadata
                    elif content_type.startswith('text/html'):
                        content, html_metadata = self._extract_html(temp_path)
                        metadata.update(html_metadata)
                        return content, metadata
                    elif content_type.startswith('text/'):
                        content, text_metadata = self._extract_text(temp_path)
                        metadata.update(text_metadata)
                        return content, metadata
                    else:
                        # Try to extract as text
                        content, text_metadata = self._extract_text(temp_path)
                        metadata.update(text_metadata)
                        return content, metadata
                finally:
                    # Clean up temporary file
                    with contextlib.suppress(Exception):
                        os.unlink(temp_path)
            
            # If request failed, return empty content with error metadata
            return "", {"error": "Failed to fetch URL", "url": url}
        except Exception as e:
            logger.error(f"Error extracting content from URL {url}: {e}")
            return "", {"error": str(e), "url": url}
    
    def _detect_mime_type(self, file_path: str) -> Optional[str]:
        """
        Attempt to detect the MIME type of a file by examining its contents.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Detected MIME type or None if unable to determine
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4096)
                
            # Check for PDF signature
            if header.startswith(b'%PDF'):
                return 'application/pdf'
            
            # Check for Office Open XML (DOCX) signature
            if header.startswith(b'PK\x03\x04'):
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Check for XML
            if header.startswith(b'<?xml'):
                return 'application/xml'
            
            # Check for HTML
            if b'<!DOCTYPE html>' in header or b'<html' in header:
                return 'text/html'
            
            # Try to detect if it's text
            try:
                encoding = chardet.detect(header)['encoding']
                if encoding:
                    header_text = header.decode(encoding)
                    
                    # Check for JSON
                    if header_text.strip().startswith('{') or header_text.strip().startswith('['):
                        try:
                            json.loads(header_text)
                            return 'application/json'
                        except:
                            pass
                    
                    # Check for CSV
                    if ',' in header_text and '\n' in header_text:
                        return 'text/csv'
                    
                    # Default to text
                    return 'text/plain'
            except:
                pass
            
            return None
        except Exception as e:
            logger.error(f"Error detecting MIME type for {file_path}: {e}")
            return None
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            text_content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        for key, value in pdf_reader.metadata.items():
                            if key.startswith('/'):
                                metadata[key[1:]] = value
                    
                    # Extract text from each page
                    num_pages = len(pdf_reader.pages)
                    metadata['num_pages'] = num_pages
                    
                    for page_num in range(num_pages):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text() or ""
                            text_content += page_text + "\n\n"
                        except Exception as page_error:
                            logger.warning(f"Error extracting text from page {page_num}: {page_error}")
                            text_content += f"[Error extracting page {page_num}]\n\n"
                except Exception as pdf_error:
                    logger.error(f"Error reading PDF: {pdf_error}")
                    return "", {"error": str(pdf_error)}
            
            return text_content.strip(), metadata
        except Exception as e:
            logger.error(f"Error extracting PDF content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Extract metadata
            metadata = {
                "core_properties": {},
                "sections": len(doc.sections),
                "paragraphs": len(doc.paragraphs)
            }
            
            # Extract core properties if available
            try:
                core_props = doc.core_properties
                for prop in dir(core_props):
                    if not prop.startswith('_') and not callable(getattr(core_props, prop)):
                        value = getattr(core_props, prop)
                        if value is not None:
                            metadata["core_properties"][prop] = str(value)
            except:
                pass
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            text_content = ""
            rows = []
            
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read(4096))
                encoding = result['encoding'] or 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                try:
                    csv_reader = csv.reader(file)
                    headers = next(csv_reader, [])
                    
                    # Add headers
                    if headers:
                        text_content += " | ".join(headers) + "\n"
                        text_content += "-" * 80 + "\n"
                    
                    # Add rows
                    row_count = 0
                    for row in csv_reader:
                        text_content += " | ".join(row) + "\n"
                        rows.append(row)
                        row_count += 1
                except Exception as csv_error:
                    logger.error(f"Error reading CSV: {csv_error}")
                    return "", {"error": str(csv_error)}
            
            metadata = {
                "headers": headers,
                "row_count": row_count,
                "encoding": encoding
            }
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting CSV content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from a JSON file.
        
        Args:
            file_path: Path to the JSON file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read(4096))
                encoding = result['encoding'] or 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                try:
                    json_data = json.load(file)
                except Exception as json_error:
                    logger.error(f"Error parsing JSON: {json_error}")
                    return "", {"error": str(json_error)}
            
            # Convert JSON to formatted string
            try:
                text_content = json.dumps(json_data, indent=2)
            except Exception as format_error:
                logger.error(f"Error formatting JSON: {format_error}")
                text_content = str(json_data)
            
            # Extract metadata
            metadata = {
                "encoding": encoding,
                "structure": "object" if isinstance(json_data, dict) else "array",
                "size": len(text_content)
            }
            
            if isinstance(json_data, dict):
                metadata["keys"] = list(json_data.keys())
            elif isinstance(json_data, list):
                metadata["items"] = len(json_data)
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting JSON content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_xml(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from an XML file.
        
        Args:
            file_path: Path to the XML file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read(4096))
                encoding = result['encoding'] or 'utf-8'
            
            # Parse XML
            try:
                tree = ET.parse(file_path)
                root = tree.getroot()
            except Exception as xml_error:
                logger.error(f"Error parsing XML: {xml_error}")
                return "", {"error": str(xml_error)}
            
            # Extract all text content
            def extract_text_from_element(element):
                text = element.text or ""
                for child in element:
                    text += extract_text_from_element(child)
                    if child.tail:
                        text += child.tail
                return text
            
            text_content = extract_text_from_element(root)
            
            # Extract metadata
            metadata = {
                "encoding": encoding,
                "root_tag": root.tag,
                "namespace": root.tag.split('}')[0].strip('{') if '}' in root.tag else None,
                "attributes": {k: v for k, v in root.attrib.items()},
                "child_elements": len(list(root))
            }
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting XML content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_html(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from an HTML file.
        
        Args:
            file_path: Path to the HTML file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                content = f.read()
                result = chardet.detect(content)
                encoding = result['encoding'] or 'utf-8'
            
            # Parse HTML
            try:
                soup = BeautifulSoup(content, 'html.parser')
            except Exception as html_error:
                logger.error(f"Error parsing HTML: {html_error}")
                return "", {"error": str(html_error)}
            
            # Extract metadata
            metadata = {
                "encoding": encoding,
                "title": soup.title.string if soup.title else None,
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img'))
            }
            
            # Extract meta tags
            meta_tags = {}
            for meta in soup.find_all('meta'):
                name = meta.get('name') or meta.get('property')
                content = meta.get('content')
                if name and content:
                    meta_tags[name] = content
            
            metadata["meta_tags"] = meta_tags
            
            # Extract text content (remove script and style elements)
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text_content = soup.get_text(separator="\n")
            
            # Clean up text (remove excessive whitespace)
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting HTML content from {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract content from a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Tuple containing extracted text and metadata
        """
        try:
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                content = f.read()
                result = chardet.detect(content)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                text_content = file.read()
            
            # Extract metadata
            metadata = {
                "encoding": encoding,
                "size": len(text_content),
                "lines": text_content.count('\n') + 1
            }
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error extracting text content from {file_path}: {e}")
            return "", {"error": str(e)}
